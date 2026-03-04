"""Drawing Agent + 措施图插入 单元测试。"""

import sys
import json
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest


# ── get_project_data 测试 ────────────────────────────────────

def test_get_project_data():
    """验证项目数据结构。"""
    from src.tools.drawing_tools import get_project_data
    from src.context import AgentContext
    from src.state import GlobalState

    state = GlobalState()
    state.Static.meta = {"project_name": "测试项目"}
    state.ETL.zones = [
        {"name": "建(构)筑物区", "area_hm2": 1.5},
        {"name": "道路广场区", "area_hm2": 0.8},
    ]
    state.Measures = [
        {"name": "排水沟C20(40×40)", "zone": "建(构)筑物区", "source": "planned"},
        {"name": "综合绿化(乔灌草)", "zone": "道路广场区", "source": "planned"},
    ]
    state.ETL.spatial_layout = {}
    state.ETL.measure_layout = []

    with AgentContext(state=state):
        # 全局查询
        result = get_project_data()
        assert result["project_name"] == "测试项目"
        assert result["total_zones"] == 2
        assert result["total_measures"] == 2
        assert len(result["zones"]) == 2

        # 分区查询
        result = get_project_data(zone_name="建(构)筑物区")
        assert result["zone"] is not None
        assert len(result["measures"]) == 1
        assert result["measures"][0]["name"] == "排水沟C20(40×40)"


# ── get_style_reference 测试 ─────────────────────────────────

def test_get_style_reference():
    """验证样式引用包含内置样式。"""
    from src.tools.drawing_tools import get_style_reference
    from src.context import AgentContext

    with AgentContext(state=None, atlas_rag=None):
        result = get_style_reference(
            map_type="zone_detail",
            measure_names=["排水沟C20(40×40)", "综合绿化(乔灌草)"],
        )

    assert "measure_styles" in result
    assert "排水沟C20(40×40)" in result["measure_styles"]
    assert result["measure_styles"]["排水沟C20(40×40)"]["type"] == "line"
    assert "zone_colors" in result
    assert "map_defaults" in result


# ── verify_image 测试 ────────────────────────────────────────

def test_verify_image_fallback():
    """VL 不可用时应降级通过 (文件足够大)。"""
    from src.tools.drawing_tools import verify_image

    with tempfile.TemporaryDirectory() as tmpdir:
        # 创建一个 > 50KB 的假图片 (Windows size_check 阈值)
        fake_img = Path(tmpdir) / "test.png"
        fake_img.write_bytes(b"\x89PNG" + b"\x00" * 60000)

        result = verify_image(str(fake_img), "zone_boundary")
        assert result["pass"] is True
        assert result["source"] in ("fallback", "size_check")
        assert result["score"] >= 70


def test_verify_image_missing():
    """不存在的文件应返回失败。"""
    from src.tools.drawing_tools import verify_image

    result = verify_image("/nonexistent/path.png", "zone_boundary")
    assert result["pass"] is False
    assert result["score"] == 0


# ── _insert_measure_maps 测试 ────────────────────────────────

def test_insert_measure_maps():
    """python-docx 图片插入应不报错。"""
    from docx import Document
    from src.renderer import _insert_measure_maps

    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)

        # 创建简单 docx
        doc = Document()
        doc.add_heading("测试报告", 0)
        doc.add_paragraph("第五章 水土保持措施设计")
        doc.add_paragraph("措施内容...")
        docx_path = tmpdir / "test.docx"
        doc.save(str(docx_path))

        # 用 matplotlib 生成真实 PNG (> 1KB)
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        fake_img = tmpdir / "zone_boundary_map.png"
        fig, ax = plt.subplots(figsize=(4, 3))
        ax.bar(["A", "B", "C"], [1, 2, 3])
        ax.set_title("Test")
        fig.savefig(str(fake_img), dpi=100)
        plt.close(fig)

        chart_paths = {"zone_boundary_map": fake_img}

        # 插入不应报错
        _insert_measure_maps(docx_path, chart_paths)

        # 验证文件增大
        doc2 = Document(str(docx_path))
        # 应该有更多段落 (原3段 + 分页+图片+标题)
        assert len(doc2.paragraphs) > 3


# ── atlas_rag 文本索引测试 ───────────────────────────────────

def test_atlas_rag_chunk_text():
    """文本分块应工作。"""
    from src.atlas_rag import AtlasRAG

    rag = AtlasRAG()
    text = """# 第一章 概述
这是第一章的内容，介绍水土保持基本概念。

## 1.1 目的
本标准规定了水土保持图的绘制方法。

# 第二章 图例
各种措施的图例样式说明。

## 2.1 排水沟
排水沟用蓝色实线表示。
"""
    chunks = rag._chunk_text(text, source_name="test")
    assert len(chunks) >= 2
    # 每个 chunk 应有 text 和 section
    for chunk in chunks:
        assert "text" in chunk
        assert "section" in chunk
        assert len(chunk["text"]) > 0


def test_atlas_rag_file_classification():
    """文件用途分类应正确。"""
    from src.atlas_rag import AtlasRAG

    rag = AtlasRAG()
    assert rag._classify_file_purpose("SL73_6-2015水利水电工程制图标准.pdf") == "制图标准"
    assert rag._classify_file_purpose("报批稿.pdf") == "范文"
    assert rag._classify_file_purpose("江苏省水土保持条例.txt") == "法规条例"
    assert rag._classify_file_purpose("GBT 50434-2018 防治标准.pdf") == "技术标准"
    assert rag._classify_file_purpose("random_file.txt") == "其他"
