"""docx 渲染器 — docxtpl 变量替换 + python-docx 手动填充循环表格。

策略:
  1. docxtpl 仅处理 {{ }} 变量替换（模板已清除所有 {%...%} 块标签）
  2. python-docx 在渲染后手动复制行填充 4 个循环表格
  3. 完整性检查: 残留标签(段落+表格) / 空表 / 文件大小(50KB~20MB)
"""

from __future__ import annotations

import copy
import re
import logging
from pathlib import Path

from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Mm
from lxml import etree

from src.state import GlobalState
from src.settings import TEMPLATE_DOCX, OUTPUT_DIR

logger = logging.getLogger(__name__)


class _DummyLoopVar:
    """占位对象，所有属性访问返回空字符串。

    docxtpl 遇到 {{ em.id }} 等循环变量时，从 context 取到此对象，
    .id 等属性返回 ""，渲染为空。后续 _fill_loop_tables() 再填入真实数据。
    """
    def __getattr__(self, name):
        return ""

    def __str__(self):
        return ""


# ── 循环表格填充 ─────────────────────────────────────────────

def _fill_loop_tables(docx_path: Path, tpl_ctx: dict) -> None:
    """打开已渲染的 docx，手动填充 4 个循环表格的数据行。

    表格索引 (基于模板结构):
      T5  — existing_measures   (em.id / em.name / em.unit / em.qty / em.location / em.cost)
      T9  — zone1_measures      (m.type / m.name / m.form / m.location / m.period / m.qty / m.unit)
      T10 — zone2_measures      (同上)
      T12 — cost_detail_table   (cd.id / cd.name / cd.unit / cd.qty / cd.price / cd.total)
    """
    doc = Document(str(docx_path))
    tables = doc.tables

    _fill_table(tables, 5, tpl_ctx.get("existing_measures", []),
                ["id", "name", "unit", "qty", "location", "cost"])
    _fill_table(tables, 9, tpl_ctx.get("zone1_measures", []),
                ["type", "name", "form", "location", "period", "qty", "unit"])
    _fill_table(tables, 10, tpl_ctx.get("zone2_measures", []),
                ["type", "name", "form", "location", "period", "qty", "unit"])
    _fill_table(tables, 12, tpl_ctx.get("cost_detail_table", []),
                ["id", "name", "unit", "qty", "price", "total"])

    doc.save(str(docx_path))


def _fill_table(tables, table_idx: int, data: list[dict], fields: list[str]) -> None:
    """用 data 填充指定表格的数据行。

    逻辑: 表格第 1 行 (index=1) 是模板行（含 {{ xx.yy }}  占位符）。
    将模板行复制 len(data) 次，每次替换占位文本为实际值。
    数据行插入到模板行原位置，保证在"合计"行之前。
    """
    if table_idx >= len(tables) or not data:
        return

    table = tables[table_idx]
    if len(table.rows) < 2:
        return

    # 获取模板行的 XML 元素
    template_row = table.rows[1]
    tbl_element = table._tbl
    template_tr = template_row._tr

    # 记录模板行在表格 XML 中的位置（用于在原位置插入数据行）
    insert_pos = list(tbl_element).index(template_tr)

    # 删除原始模板行（里面是 {{ xx.yy }} 占位）
    tbl_element.remove(template_tr)

    # 为每条数据复制一行，在原模板行位置依次插入
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for i, item in enumerate(data):
        new_tr = copy.deepcopy(template_tr)
        # 遍历 <w:tc> 单元格
        cells = new_tr.findall(".//w:tc", ns)
        for ci, field in enumerate(fields):
            if ci < len(cells):
                value = str(item.get(field, ""))
                # 替换单元格内所有 <w:t> 的文本
                t_nodes = cells[ci].findall(".//w:t", ns)
                if t_nodes:
                    t_nodes[0].text = value
                    for extra in t_nodes[1:]:
                        extra.text = ""
        tbl_element.insert(insert_pos + i, new_tr)


# ── 主渲染函数 ──────────────────────────────────────────────

def render_docx(state: GlobalState, output_path: Path | None = None,
                chart_paths: dict[str, Path] | None = None) -> Path:
    """渲染 docx 报告。"""
    if output_path is None:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        output_path = OUTPUT_DIR / "report.docx"

    tpl_ctx = dict(state.TplCtx)

    tpl = DocxTemplate(str(TEMPLATE_DOCX))

    # 注入图表 InlineImage
    if chart_paths:
        for name, path in chart_paths.items():
            if path.exists():
                try:
                    tpl_ctx[name] = InlineImage(tpl, str(path), width=Mm(150))
                except Exception as e:
                    logger.warning(f"图片 {name} 注入失败: {e}")

    # 注入循环变量占位对象，避免 docxtpl 遇到 {{ em.id }} 等报错
    for loop_var in ("em", "m", "cd"):
        if loop_var not in tpl_ctx:
            tpl_ctx[loop_var] = _DummyLoopVar()

    # Step 1: docxtpl 渲染变量替换（模板已无 {%...%} 块标签）
    try:
        tpl.render(tpl_ctx)
        tpl.save(str(output_path))
    except Exception as e:
        logger.error(f"docxtpl 渲染失败: {e}")
        raise

    # Step 2: python-docx 手动填充循环表格
    try:
        _fill_loop_tables(output_path, tpl_ctx)
    except Exception as e:
        logger.warning(f"循环表格填充失败(非致命): {e}")

    # Step 3: 后插入措施图 (python-docx)
    if chart_paths:
        try:
            _insert_measure_maps(output_path, chart_paths)
        except Exception as e:
            logger.warning(f"措施图插入失败(非致命): {e}")

    logger.info(f"报告已生成: {output_path}")
    return output_path


# ── 措施图后插入 ────────────────────────────────────────────

# 措施图排序权重 (数值越小越靠前)
_MAP_ORDER = {
    "zone_boundary_map": (1, "图5-1 水土保持防治分区图"),
    "measure_layout_map": (2, "图5-2 水土保持措施总体布置图"),
}


def _insert_measure_maps(docx_path: Path, chart_paths: dict[str, Path]) -> None:
    """后置插入措施图 (python-docx)。

    在第五章末尾按顺序插入措施图 (而非追加到文档最末):
      1. zone_boundary_map — 图5-1 水土保持防治分区图
      2. measure_layout_map — 图5-2 水土保持措施总体布置图
      3. zone_detail_* × N — 图5-3~5-N 各分区措施详图
      4. typical_section_* × M — 图5-(N+1)~5-(N+M) 典型断面图
    """
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    # 筛选措施图 (排除 4 张固定数据图表，名称须与 charts.py 一致)
    fixed_charts = {"erosion_chart", "investment_pie", "benefit_bar", "zone_pie"}
    measure_maps = {
        name: path for name, path in chart_paths.items()
        if name not in fixed_charts and path.exists() and path.stat().st_size > 1024
    }

    if not measure_maps:
        logger.info("  无有效措施图可插入")
        return

    # 排序
    sorted_maps = _sort_measure_maps(measure_maps)

    # 打开文档
    doc = Document(str(docx_path))

    # 查找插入位置: 第五章末尾段落
    insert_after_idx = _find_chapter5_end(doc)

    # 获取插入参考点 (第五章末段的 XML 元素)
    ref_element = doc.paragraphs[insert_after_idx]._element
    parent = ref_element.getparent()
    insert_pos = list(parent).index(ref_element) + 1

    # 逐张插入 (每次 insert_pos 递增)
    fig_num = 1
    for tag, path, title in sorted_maps:
        # 1. 创建分页段落
        page_break = doc.add_paragraph()
        page_break.paragraph_format.space_before = Pt(0)
        page_break.paragraph_format.space_after = Pt(0)
        run = page_break.add_run()
        br = run._element.makeelement(qn("w:br"), {qn("w:type"): "page"})
        run._element.append(br)
        # 从末尾移到正确位置
        pb_elem = page_break._element
        parent.remove(pb_elem)
        parent.insert(insert_pos, pb_elem)
        insert_pos += 1

        # 2. 创建图片段落
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic_para.add_run()
        try:
            run.add_picture(str(path), width=Cm(15))
        except Exception as e:
            logger.warning(f"  插入图片失败 {tag}: {e}")
            # 移除空段落
            parent.remove(pic_para._element)
            continue
        # 从末尾移到正确位置
        pic_elem = pic_para._element
        parent.remove(pic_elem)
        parent.insert(insert_pos, pic_elem)
        insert_pos += 1

        # 3. 创建标题段落
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_before = Pt(6)
        caption.paragraph_format.space_after = Pt(12)
        cap_run = caption.add_run(f"图5-{fig_num} {title}")
        cap_run.font.size = Pt(10.5)
        # 从末尾移到正确位置
        cap_elem = caption._element
        parent.remove(cap_elem)
        parent.insert(insert_pos, cap_elem)
        insert_pos += 1

        fig_num += 1

    doc.save(str(docx_path))
    logger.info(f"  措施图插入: {len(sorted_maps)} 张 (位于第五章后)")


def _sort_measure_maps(maps: dict[str, Path]) -> list[tuple[str, Path, str]]:
    """按规则排序措施图，返回 [(tag, path, title), ...]。"""
    sorted_list = []

    # 1. 固定图 (分区图 + 总布置图)
    for tag in ("zone_boundary_map", "measure_layout_map"):
        if tag in maps:
            _, title = _MAP_ORDER.get(tag, (99, tag))
            sorted_list.append((tag, maps[tag], title))

    # 2. 分区详图 (zone_detail_*)
    detail_maps = sorted(
        [(k, v) for k, v in maps.items() if k.startswith("zone_detail_")],
        key=lambda x: x[0],
    )
    for tag, path in detail_maps:
        zone_name = tag.replace("zone_detail_", "")
        sorted_list.append((tag, path, f"{zone_name}措施详图"))

    # 3. 断面图 (typical_section_*)
    section_maps = sorted(
        [(k, v) for k, v in maps.items() if k.startswith("typical_section_")],
        key=lambda x: x[0],
    )
    for tag, path in section_maps:
        section_name = tag.replace("typical_section_", "")
        sorted_list.append((tag, path, f"{section_name}典型断面图"))

    return sorted_list


def _find_chapter5_end(doc: Document) -> int:
    """查找第五章末尾的段落索引，用于定位插入位置。"""
    chapter5_found = False
    last_ch5_idx = len(doc.paragraphs) - 1

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # 匹配 "第五章" 或 "5 措施设计" 或 "第5章"
        if re.search(r"(第五章|第5章|^5[\s\.．])", text):
            chapter5_found = True
        # 找到第五章后，如果遇到下一章则返回前一个位置
        if chapter5_found and re.search(r"(第六章|第6章|^6[\s\.．])", text):
            return i - 1

    return last_ch5_idx


# ── 完整性检查 ──────────────────────────────────────────────

def check_completeness(docx_path: Path) -> dict:
    """检查渲染后 docx 的完整性。

    检查项:
      1. 文件大小: 50KB ~ 20MB (PRD §8.3)
      2. 残留 {{ }} 或 {%...%} 标签 (段落+表格单元格)
      3. 全空表格
      4. 非空段落占比
    """
    issues = []

    size_kb = docx_path.stat().st_size / 1024
    if size_kb < 50:
        issues.append(f"文件过小: {size_kb:.1f}KB (要求≥50KB)")
    if size_kb > 20480:
        issues.append(f"文件过大: {size_kb:.1f}KB (要求≤20MB)")

    try:
        doc = Document(str(docx_path))

        # 残留 {{ }} 标签 — 同时扫描段落和表格单元格
        residual_tags = []
        for para in doc.paragraphs:
            text = para.text
            tags = re.findall(r"\{\{.*?\}\}", text)
            residual_tags.extend(tags)
            block_tags = re.findall(r"\{%.*?%\}", text)
            residual_tags.extend(block_tags)

        # 扫描表格单元格中的残留标签
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text
                    tags = re.findall(r"\{\{.*?\}\}", text)
                    residual_tags.extend(tags)
                    block_tags = re.findall(r"\{%.*?%\}", text)
                    residual_tags.extend(block_tags)

        if residual_tags:
            issues.append(f"残留标签 ({len(residual_tags)}个): {residual_tags[:5]}")

        # 空表格
        empty_tables = 0
        for table in doc.tables:
            cell_texts = [cell.text.strip() for row in table.rows for cell in row.cells]
            if all(not t for t in cell_texts):
                empty_tables += 1
        if empty_tables:
            issues.append(f"空表格: {empty_tables}个")

        # 空段落占比
        total_paras = len(doc.paragraphs)
        non_empty = sum(1 for p in doc.paragraphs if p.text.strip())
        if total_paras > 0 and non_empty / total_paras < 0.3:
            issues.append(f"非空段落比例过低: {non_empty}/{total_paras}")

    except Exception as e:
        issues.append(f"检查异常: {str(e)}")

    return {
        "file_size_kb": round(size_kb, 1),
        "issues": issues,
        "pass": len(issues) == 0,
    }
